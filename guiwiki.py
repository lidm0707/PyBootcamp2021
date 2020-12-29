#GUIWiki.py
import wikipedia

#python to doc x
from docx import Document

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #summary สำหรีบบทความที่สรุป
    data = wikipedia.summary(keyword)

    doc=Document() # สร้างไฟล์ word ใน python
    doc.add_heading(keyword , 0)

    #page + content ฐทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('FINISH')

#เปลี่นรเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI= Tk()
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')
#config
FONT1 = ('Angsana New',15)

# คำอธิบาย
L = ttk.Label(GUI, text='ค้นหาบทความ',font=FONT1)
L.pack()


# ช่องต้นหาข้อมูล
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

# ปุ่มค้นหา
def Search():
    keyword = v_search.get() # .get คือ ดึงข้แมูล
    try:
        # try to srearch. success is padd.
        language = v_radio.get()
        Wiki(keyword,language)

        messagebox.showinfo('save','Search is success saved')
        # if promble is pass
    except:
        messagebox.showwarning('Keyword Error','Try again')

    # print(wikipedia.search(keyword))
    # result = wikipedia.summary(keyword )
    # print(result)
    
B1 = ttk.Button(GUI,text='search',command=Search)
B1.pack(ipadx=20,ipady=10)

# menu select language
F1 = Frame(GUI)
F1.pack(pady=10,padx=10)

v_radio = StringVar()

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย' , variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='eng' , variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='chi' , variable=v_radio,value='zh')
RB1.invoke() # default thai

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()
